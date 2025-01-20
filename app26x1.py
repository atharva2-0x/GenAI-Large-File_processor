import streamlit as st
from fiservai import FiservAI
from dotenv import load_dotenv
import os
import time
import asyncio
import re
from docx import Document
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
from threading import Lock
from constants import *

load_dotenv()

API_KEY = os.getenv("API_KEY")
if API_KEY is None or API_KEY == "":
    raise ValueError("API_KEY is not set. Check your environment variables or add a .env file with API_KEY.")

API_SECRET = os.getenv("API_SECRET")
if API_SECRET is None or API_SECRET == "":
    raise ValueError("API_SECRET is not set. Check your environment variables or add a .env file with API_SECRET.")


BASE_URL = os.getenv("BASE_URL")

CONVERSATION_DIR = 'conversation'
RESPONSE_DIR = 'responses'

for directory in [CONVERSATION_DIR, RESPONSE_DIR]:
    if not os.path.exists(directory):
        os.makedirs(directory)

client = FiservAI.FiservAI(API_KEY, API_SECRET, base_url=BASE_URL if BASE_URL is not None else None)

def initialize_session_state():
    if "prompt" not in st.session_state:
        st.session_state.prompt = []
    if "response" not in st.session_state:
        st.session_state.response = []
    if "new_conversation" not in st.session_state:
        st.session_state.new_conversation = True

initialize_session_state()

def print_diagnostics(resp, elapsed_time_sec):
    st.write(f"""[Model: {resp.model}, Prompt Tokens: {resp.usage.prompt_tokens}, Completion Tokens: {resp.usage.completion_tokens}, Total Tokens: {resp.usage.total_tokens}, Cost: ${resp.usage.cost}, Time: {elapsed_time_sec} sec]""")

def stream_headline(response):
    for word in response:
        yield word + ""
        time.sleep(0.01)

def stream_line(response):
    for line in response.splitlines():
        yield line + "\n"
        time.sleep(0.1)

def get_conversation_file(task_option, uploaded_file_name):
    if 'conversation_file' not in st.session_state or st.session_state.new_conversation:
        if uploaded_file_name:
            base_name = os.path.splitext(uploaded_file_name)[0]
        else:
            base_name = "conversation"
        
        sanitized_task_option = "".join([c if c.isalnum() else "_" for c in task_option]).strip("_")
        st.session_state.conversation_file = os.path.join(
            CONVERSATION_DIR, 
            f'{sanitized_task_option}_{base_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt'
        )
        st.session_state.new_conversation = False
    return st.session_state.conversation_file



file_lock = Lock()

def save_user_convo(user_prompt, task_option, uploaded_file_name):
    with file_lock:
        with open(get_conversation_file(task_option, uploaded_file_name), "a") as f:
            f.write(f"\n\n\n\n\nuser: {user_prompt}")

def save_assist_convo(assistant_response, task_option, uploaded_file_name):
    with file_lock:        
        with open(get_conversation_file(task_option, uploaded_file_name), "a") as f:
            f.write(f"\n\n\n\n\nassistant: {assistant_response}")


def remove_triple_quotes(text):
    lines = text.split('\n')
    filtered_lines = [line for line in lines if "```" not in line]
    return '\n'.join(filtered_lines)

def download_last_response():
    if st.session_state.response:
        response = st.session_state.response[-1]
        response = remove_triple_quotes(response)  
        doc = Document()
        doc.add_heading('Assistant Response', level=1)
        doc.add_paragraph(response)
        file_path = os.path.join(RESPONSE_DIR, f'assistant_response_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
        doc.save(file_path)
        with open(file_path, 'rb') as f:
            st.download_button('Download Last Response ⬇️', f, file_name=f'assistant_response_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')

def download_entire_response(entire_response):
    entire_response = remove_triple_quotes(entire_response)  
    doc = Document()
    doc.add_heading('Assistant Response', level=1)
    doc.add_paragraph(entire_response)
    file_path = os.path.join(RESPONSE_DIR, f'entire_assistant_response_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    doc.save(file_path)
    with open(file_path, 'rb') as f:
        st.download_button('Download Entire Response ⬇️', f, file_name=f'entire_assistant_response_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')

def start_new_conversation():
    st.session_state.prompt = []
    st.session_state.response = []
    st.session_state.new_conversation = True



def load_conversation(file_path):
    # Open the conversation file and read its content
    with open(file_path, "r") as f:
        content = f.read()

    # Split content by lines and parse messages
    conversation_lines = content.split('\n\n\n\n\n')
    for line in conversation_lines:
        if line.startswith("user:"):
            text = line.split('\n')[1:]
            st.session_state.prompt.append(''.join(text).strip())
        elif line.startswith("assistant:"):
            # Append assistant response to session state response list
            st.session_state.response.append(line[10:].strip())




# Function to extract datetime from filename
def get_datetime_from_filename(filename):
    match = re.search(r"_(\d{8}_\d{6})", filename)
    if match:
        return datetime.strptime(match.group(1), "%Y%m%d_%H%M%S")
    return datetime.min



def chunk_text_by_lines(full_prompt, task_option, first_chunk_size=100, subsequent_chunk_size=150):
    lines = full_prompt.split('\n')
    chunks = []
    chunk = []
    current_chunk_size = first_chunk_size

    for line in lines:
        if line.strip() == '' and len(chunk) >= current_chunk_size:
            if chunk:
                chunks.append("\n".join(chunk))
                chunk = []
                current_chunk_size = subsequent_chunk_size  
        else:
            chunk.append(line)

    if chunk:
        chunks.append("\n".join(chunk))

    chunks = [task_option + "\n\n" + chunk for chunk in chunks]
    return chunks


async def fetch_response(client, messages):
            response = await client.chat_completion_async(messages)
            response_content = response.choices[0].message.content
            return response_content


def fetch_response_sync(messages):
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    response = loop.run_until_complete(client.chat_completion_async(messages))
    return response.choices[0].message.content


def process_chunks(chunks):
    responses = []
    with ThreadPoolExecutor() as executor:
        futures = []
        for chunk in chunks:
            st.session_state.prompt.append(chunk)
            save_user_convo(chunk, task_option, uploaded_file_name)
            messages = [
                {"role": "user", "content": m}
                for m in st.session_state.prompt[-3:]
            ] + [
                {"role": "assistant", "content": m}
                for m in st.session_state.response[-3:]
            ]
            
            futures.append(executor.submit(fetch_response_sync, messages))
        
        for future in futures:
            response_content = future.result()
            responses.append(response_content)
            st.session_state.response.append(response_content)
            save_assist_convo(response_content, task_option, uploaded_file_name)
            st.write_stream(stream_line(response_content))
    return responses


def main(full_prompt, task_prompt):
    entire_response = ""
    chunks = chunk_text_by_lines(full_prompt, task_prompt)
    entire_response_list =process_chunks(chunks)
    entire_response= "\n".join(entire_response_list)

    col1, col2 = st.columns(2)
    with col1:
        download_last_response()
    with col2:
        download_entire_response(entire_response)


def inject_custom_css():
    st.markdown(
        """
        <style>
        .custom-background {
            background-color: #ff6803 !important;
            padding: 10px;
            border-radius: 5px;
        }
        .stButton {
            color: #ff6803 !important;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

inject_custom_css()

def delete_file(file_path):
    if os.path.exists(file_path):
        os.remove(file_path)
        time.sleep(0.5)
        st.sidebar.error(f"Deleted ❌")
        st.rerun()
        
    

st.sidebar.title("Choose a Task ⚒️")

task_option = st.sidebar.selectbox(
    "Choose a task",
    (
        "Enter a custom prompt",
        "Generate Documentation",
        "Adding Comments",
        "Analyse the program",
        "Generate Test cases",
        "Generate A Scrum Story",
    ),
)

tasks_prompts = {
    "Enter a custom prompt": CUSTOM_PROMPT,  
    "Generate Documentation": DOCGEN,
    "Adding Comments": COMMENT_GEN,
    "Analyse the program": ANALYZER,
    "Generate Test cases": GEN_T_CASE,
    "Generate A Scrum Story": SCRUM_STORY,
}

st.sidebar.write("")

if st.sidebar.button('Start New Conversation'):
    start_new_conversation()

filetext = ""
uploaded_file_name = ""
uploaded_file = st.sidebar.file_uploader("Drag & drop files here to upload file.", type=None)
if uploaded_file is not None:
    filetext = uploaded_file.read().decode("utf-8")
    uploaded_file_name = uploaded_file.name

st.sidebar.title("Conversation History 📃")
conversation_files = os.listdir(CONVERSATION_DIR)
conversation_files = sorted(conversation_files, key=get_datetime_from_filename, reverse=True)
for file in conversation_files:
    col1, col2 = st.sidebar.columns([4, 1])
    with col1:
        if st.button(file):
            start_new_conversation()
            load_conversation(os.path.join("conversation", file))
    with col2:
        if st.button("✕", key=f"delete_{file}"):
            start_new_conversation()
            delete_file(os.path.join(CONVERSATION_DIR, file))


for prompt, response in zip(st.session_state.prompt, st.session_state.response):
    with st.chat_message("user"):
        st.markdown(prompt)
    with st.chat_message("assistant"):
        st.markdown(response)


headline = "Welcome, How can I help? 😎"
if task_option == "Generate Documentation":
    headline = "Docgen 📜"
elif task_option == "Adding Comments":
    headline = "Coments Adder 💬"
elif task_option == "Analyse the program":
    headline = "Program Analyzer 🔍"
elif task_option == "Generate Test cases":
    headline = "Test Case Generator 🧪"
elif task_option == "Generate A Scrum Story":
    headline = "Story Generator 📝"

st.write_stream(stream_headline(f"# **{headline}**"))
st.markdown('<div class="custom-background">', unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)

if prompt := st.chat_input("Write a prompt or Choose a Task for me 🤖"):
    full_prompt = f" {prompt} \t {filetext}"
    full_prompt_lines = full_prompt.split('\n')
    
    
    with st.chat_message("user"):
        st.markdown(prompt)

    if len(full_prompt_lines) > 200:
        with st.chat_message("assistant"):
            main(full_prompt, tasks_prompts[task_option])
       
    else:
        entire_response =""
        st.session_state.prompt.append(tasks_prompts[task_option] + full_prompt)
        save_user_convo(tasks_prompts[task_option] + full_prompt, task_option, uploaded_file_name)

        with st.chat_message("assistant"):

            messages = [
                {"role": "user", "content": m}
                for m in st.session_state.prompt[-4:]
            ] + [
                {"role": "assistant", "content": m}
                for m in st.session_state.response[-4:]
            ]
            
            response_stream = asyncio.run(fetch_response(client, messages))
            entire_response += response_stream
            st.write_stream(stream_line(response_stream))
            st.session_state.response.append(response_stream)

            save_assist_convo(response_stream, task_option, uploaded_file_name)
            col1, col2 = st.columns(2)
            with col1:
                download_last_response()
            with col2:
                download_entire_response(entire_response)

        
            